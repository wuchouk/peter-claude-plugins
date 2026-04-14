#!/usr/bin/env python3
"""fill-docx.py — Word 文件填寫：解密資料 + python-docx 填寫

用法: fill-docx.py <docx_file> <mapping.json>

支援兩種模式：
  - table: 填寫表格中的指定儲存格
  - placeholder: 替換 {{placeholder}} 或 ____ 佔位符

⚠️ stdout 只輸出狀態訊息，不輸出任何資料值
"""

import json
import os
import re
import subprocess
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt


def resolve_key(data, key):
    """解析 dot notation key，支援陣列 [index]"""
    parts = re.split(r'\.|\[(\d+)\]', key)
    parts = [p for p in parts if p is not None and p != '']
    current = data
    for part in parts:
        if part.isdigit():
            idx = int(part)
            if isinstance(current, list) and idx < len(current):
                current = current[idx]
            else:
                return None
        elif isinstance(current, dict) and part in current:
            current = current[part]
        else:
            return None
    return current


def apply_transform(value, transform):
    """套用值轉換"""
    if not transform or value is None:
        return value

    value = str(value)

    if transform.startswith("date_format:"):
        fmt = transform.split(":", 1)[1]
        try:
            dt = datetime.strptime(value, "%Y-%m-%d")
            result = fmt
            result = result.replace("YYYY", str(dt.year))
            result = result.replace("YY", str(dt.year)[-2:])
            result = result.replace("MM", f"{dt.month:02d}")
            result = result.replace("DD", f"{dt.day:02d}")
            result = result.replace("month_name", dt.strftime("%B"))
            result = result.replace("month_abbr", dt.strftime("%b"))
            return result
        except ValueError:
            return value

    if transform.startswith("date_part:"):
        part = transform.split(":", 1)[1]
        try:
            dt = datetime.strptime(value, "%Y-%m-%d")
            parts_map = {
                "year": str(dt.year),
                "month": f"{dt.month:02d}",
                "day": f"{dt.day:02d}",
                "month_name": dt.strftime("%B"),
                "month_abbr": dt.strftime("%b"),
            }
            return parts_map.get(part, value)
        except ValueError:
            return value

    if transform.startswith("select_map:"):
        map_str = transform.split(":", 1)[1].strip("{}")
        mapping = {}
        for pair in map_str.split(","):
            k, v = pair.split(":", 1)
            mapping[k.strip()] = v.strip()
        return mapping.get(value, mapping.get("*", value))

    if transform == "uppercase":
        return value.upper()
    if transform == "lowercase":
        return value.lower()
    if transform == "phone_format":
        return re.sub(r'[\s\-\(\)]', '', value)
    if transform == "phone_international":
        match = re.match(r'(\+\d{1,3})(0+)(.*)', value)
        if match:
            return match.group(1) + match.group(3)
        return value
    if transform.startswith("truncate:"):
        n = int(transform.split(":", 1)[1])
        return value[:n]

    return value


def decrypt_data(script_dir):
    """呼叫 encrypt.sh 解密，回傳暫存檔路徑（金鑰從 macOS Keychain 取得）"""
    cmd = ["bash", os.path.join(script_dir, "encrypt.sh"), "--decrypt"]

    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"❌ 解密失敗: {result.stderr.strip()}", file=sys.stderr)
        sys.exit(1)

    tmpfile = result.stdout.strip()
    if not os.path.exists(tmpfile):
        print("❌ 解密暫存檔不存在", file=sys.stderr)
        sys.exit(1)

    return tmpfile


def replace_in_paragraph(paragraph, old_text, new_text):
    """替換段落中的文字，保留原始格式"""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # 嘗試在單一 run 中替換
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True

    # 如果 old_text 跨越多個 run，需要合併處理
    # 保留第一個 run 的格式，清除其他 run
    new_full = full_text.replace(old_text, new_text)
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = new_full
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def fill_docx(docx_path, mapping_path):
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # 解密
    print("🔐 解密個人資料...")
    tmpfile = decrypt_data(script_dir)

    try:
        with open(tmpfile, "r") as f:
            data = yaml.safe_load(f)

        with open(mapping_path, "r") as f:
            mappings = json.load(f)

        # 讀取 Word
        print(f"📝 讀取 Word: {docx_path}")
        doc = Document(docx_path)

        filled = 0
        errors = []

        for m in mappings:
            fill_type = m["type"]
            key = m.get("key", "")
            transform = m.get("transform", "")

            if "keys" in m:
                values = [resolve_key(data, k) for k in m["keys"]]
                values = [str(v) for v in values if v is not None]
                if not values:
                    errors.append(f"資料 keys 不存在: {m['keys']}")
                    continue
                value = m.get("join", " ").join(values)
            else:
                value = resolve_key(data, key)
                if value is None:
                    errors.append(f"資料 key 不存在: {key}")
                    continue

            value = apply_transform(value, transform)
            if value is None:
                continue

            value = str(value)

            if fill_type == "table":
                # 表格模式
                table_idx = m["table_idx"]
                row = m["row"]
                col = m["col"]

                if table_idx >= len(doc.tables):
                    errors.append(f"表格 index {table_idx} 不存在（共 {len(doc.tables)} 個）")
                    continue

                table = doc.tables[table_idx]
                if row >= len(table.rows):
                    errors.append(f"表格 {table_idx} 的 row {row} 不存在")
                    continue
                if col >= len(table.rows[row].cells):
                    errors.append(f"表格 {table_idx} 的 col {col} 不存在")
                    continue

                cell = table.rows[row].cells[col]
                # 保留格式，清除內容後填入
                if cell.paragraphs:
                    for para in cell.paragraphs:
                        if para.runs:
                            para.runs[0].text = value
                            for run in para.runs[1:]:
                                run.text = ""
                            break
                    else:
                        cell.paragraphs[0].text = value
                else:
                    cell.text = value
                filled += 1

            elif fill_type == "placeholder":
                # 佔位符模式
                placeholder_text = m["text"]
                found = False

                # 搜尋段落
                for paragraph in doc.paragraphs:
                    if placeholder_text in paragraph.text:
                        if replace_in_paragraph(paragraph, placeholder_text, value):
                            found = True

                # 搜尋表格中的段落
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder_text in paragraph.text:
                                    if replace_in_paragraph(paragraph, placeholder_text, value):
                                        found = True

                if found:
                    filled += 1
                else:
                    errors.append(f"佔位符未找到: {placeholder_text}")

            else:
                errors.append(f"未知的填寫類型: {fill_type}")

        # 輸出
        output_path = docx_path.replace(".docx", "_filled.docx")
        doc.save(output_path)

        print(f"✅ 填入 {filled} 個欄位")
        if errors:
            print(f"⚠️  {len(errors)} 個錯誤:")
            for e in errors:
                print(f"   - {e}")
        print(f"📁 輸出: {output_path}")

    finally:
        # 安全刪除暫存檔
        if os.path.exists(tmpfile):
            subprocess.run(["rm", "-P", tmpfile], capture_output=True)
            if os.path.exists(tmpfile):
                os.remove(tmpfile)
            print("🗑️  暫存檔已清除")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: fill-docx.py <docx_file> <mapping.json>", file=sys.stderr)
        sys.exit(1)

    docx_file = sys.argv[1]
    mapping_file = sys.argv[2]
    if not os.path.exists(docx_file):
        print(f"❌ 找不到 Word 文件: {docx_file}", file=sys.stderr)
        sys.exit(1)
    if not os.path.exists(mapping_file):
        print(f"❌ 找不到映射檔: {mapping_file}", file=sys.stderr)
        sys.exit(1)

    fill_docx(docx_file, mapping_file)
